"""
parser.py — Extracts raw text from resume files (.pdf, .docx, .txt)
"""
import io
import os
import re


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extract raw text from an uploaded file in memory."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return content.decode("utf-8", errors="ignore")

    if ext == ".pdf":
        try:
            import PyPDF2
        except ImportError:
            raise RuntimeError("PyPDF2 is required to read PDF files.")
        text = []
        with io.BytesIO(content) as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)

    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx is required to read DOCX files.")
        d = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in d.paragraphs)

    raise ValueError(f"Unsupported file type: {ext}")


def extract_text(filepath: str) -> str:
    """Extract raw text from a resume file, dispatching by extension."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    elif ext == ".pdf":
        try:
            import PyPDF2
        except ImportError:
            raise RuntimeError("PyPDF2 is required to read PDF files.")
        text = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)

    elif ext == ".docx":
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx is required to read DOCX files.")
        d = docx.Document(filepath)
        return "\n".join(p.text for p in d.paragraphs)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def clean_text(text: str) -> str:
    """Normalize whitespace and strip non-informative characters."""
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s\.\+\#/\-]", " ", text)
    return text.strip()


def load_resumes_from_dir(dir_path: str) -> dict:
    """Load and clean every resume in a directory. Returns {filename: text}."""
    resumes = {}
    for fname in sorted(os.listdir(dir_path)):
        fpath = os.path.join(dir_path, fname)
        if os.path.isfile(fpath) and fname.lower().endswith((".txt", ".pdf", ".docx")):
            try:
                raw = extract_text(fpath)
                resumes[fname] = clean_text(raw)
            except Exception as e:
                print(f"[warn] could not read {fname}: {e}")
    return resumes
